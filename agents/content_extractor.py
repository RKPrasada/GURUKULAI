"""
NAGA content extractor — turns an uploaded PDF/DOCX for one subtopic into
structured study notes + practice questions.

Flow:
  1. Extract raw text from the uploaded file (pypdf / python-docx).
  2. Ask the LLM to split it into (a) clean Markdown notes and (b) MCQs.
  3. Caller persists notes (→ knowledge base) and questions (→ practice bank).

The LLM step is the only unreliable part (free-tier models are often rate-
limited). This module surfaces failure clearly so the caller can offer a Retry
button — extraction is a one-off NAGA action and is not time-critical.
"""

from __future__ import annotations

import json
import logging
import re

from agents.base import call_gemini

logger = logging.getLogger(__name__)

EXTRACT_TIMEOUT = 40.0   # allow a bit longer — documents are large

SYSTEM_PROMPT = """You are an expert content editor for Indian competitive exams.
You are given the raw text of a study document for ONE subtopic.
Split it into two parts and return ONLY valid JSON (no markdown fences):

{
  "notes": "<clean, well-structured Markdown study notes>",
  "questions": [
    {
      "question_text": "...",
      "options": ["A", "B", "C", "D"],
      "correct_index": 0,
      "explanation": "...",
      "difficulty": 1
    }
  ]
}

Rules:
- "notes": rewrite the study material as clear Markdown with ## headings,
  bullet points, **bold** key terms, and readable formulae in plain text
  (e.g. 'Mean = Sum / Count'). Include solved examples if present. Do NOT
  invent content that is not supported by the document.
- "questions": extract every MCQ present in the document. Each needs exactly
  4 options, a 0-based correct_index, a concise explanation, and difficulty
  (1=easy, 2=medium, 3=hard). If the document has no questions, return [].
- If the document has only questions and no notes, set "notes" to "".
- Return ONLY the JSON object."""


class ExtractionError(Exception):
    """Raised when text extraction or LLM parsing fails (caller offers Retry)."""


def extract_text_from_file(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded PDF or DOCX file."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="replace")
    raise ExtractionError(f"Unsupported file type: {filename!r}. Upload a PDF, DOCX, TXT, or MD file.")


def _extract_pdf(data: bytes) -> str:
    import io
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(parts).strip()
    except Exception as e:
        raise ExtractionError(f"Could not read PDF: {e}")
    if not text:
        raise ExtractionError("The PDF has no extractable text (it may be a scanned image).")
    return text


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table cells (question banks are often tabular)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
    except Exception as e:
        raise ExtractionError(f"Could not read DOCX: {e}")
    if not text:
        raise ExtractionError("The DOCX file has no readable text.")
    return text


def _strip_fences(raw: str) -> str:
    clean = raw.strip()
    for fence in ("```json", "```"):
        if clean.startswith(fence):
            clean = clean[len(fence):]
    if clean.endswith("```"):
        clean = clean[:-3]
    return clean.strip()


def extract_structured(
    text: str, exam: str, subject: str, topic: str, subtopic: str,
) -> dict:
    """
    Ask the LLM to split document text into {notes, questions}.
    Raises ExtractionError if the LLM is unavailable or returns unusable output
    (so the caller can present a Retry button).
    """
    if not text.strip():
        raise ExtractionError("No text to process.")

    # Cap very large documents to keep within model context limits
    body = text[:24000]
    prompt = (
        f"Exam: {exam.upper().replace('_', ' ')}\n"
        f"Subject: {subject}\nTopic: {topic}\nSubtopic: {subtopic}\n\n"
        f"Document text:\n{body}\n\n"
        f"Split into notes + questions per the system instructions. Return ONLY JSON."
    )
    raw = call_gemini(prompt, SYSTEM_PROMPT, timeout=EXTRACT_TIMEOUT)
    if not raw:
        raise ExtractionError(
            "The AI model is unavailable right now (likely rate-limited). "
            "Please click Retry in a moment."
        )

    try:
        data = json.loads(_strip_fences(raw))
    except Exception:
        # Best-effort salvage: find the first {...} block
        m = re.search(r"\{.*\}", _strip_fences(raw), re.DOTALL)
        if not m:
            raise ExtractionError("The AI returned an unreadable response. Please Retry.")
        try:
            data = json.loads(m.group(0))
        except Exception:
            raise ExtractionError("The AI returned malformed content. Please Retry.")

    notes = (data.get("notes") or "").strip()
    questions = data.get("questions") or []

    # Normalise and validate questions
    clean_qs: list[dict] = []
    for q in questions:
        opts = q.get("options") or []
        if not q.get("question_text") or len(opts) != 4:
            continue
        try:
            ci = int(q.get("correct_index", 0))
        except (TypeError, ValueError):
            ci = 0
        clean_qs.append({
            "question_text": str(q["question_text"]).strip(),
            "options": [str(o).strip() for o in opts],
            "correct_index": max(0, min(3, ci)),
            "explanation": str(q.get("explanation", "")).strip(),
            "difficulty": int(q.get("difficulty", 1)) if str(q.get("difficulty", 1)).isdigit() else 1,
            "subject": subject,
            "topic": topic,
            "subtopic": subtopic,
        })

    if not notes and not clean_qs:
        raise ExtractionError(
            "Could not extract any notes or questions from this document. "
            "Check the file has readable text, then Retry."
        )

    return {"notes": notes, "questions": clean_qs}
